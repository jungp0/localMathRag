from __future__ import annotations

import csv
import hashlib
import json
from pathlib import Path
from typing import Iterable

from .config import ParserConfig
from .formula import extract_formula_candidates, stable_id
from .models import ParsedBlock, ParsedDocument, SourceRef
from .ocr import try_ocr_image
from .visual import (
    build_visual_index_text,
    classify_visual_type,
    compact_text,
    copy_local_visual,
    extract_markdown_visuals,
    image_size,
    is_caption_line,
    nearest_caption,
    save_visual_artifact,
)


PARSER_VERSION = "lookup-parser-v1"
SUPPORTED_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".jsonl",
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".bmp",
    ".tif",
    ".tiff",
    ".webp",
    ".svg",
}


class DocumentParser:
    def __init__(self, config: ParserConfig | None = None):
        self.config = config or ParserConfig()

    def parse_path(self, path: str | Path, recursive: bool = True) -> list[ParsedDocument]:
        target = Path(path)
        if target.is_dir():
            files = self.iter_supported_files(target, recursive=recursive)
            return [self.parse_file(item) for item in files]
        return [self.parse_file(target)]

    def iter_supported_files(self, root: Path, recursive: bool = True) -> list[Path]:
        pattern = "**/*" if recursive else "*"
        return sorted(
            item
            for item in root.glob(pattern)
            if item.is_file() and item.suffix.lower() in SUPPORTED_SUFFIXES
        )

    def parse_file(self, path: str | Path) -> ParsedDocument:
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(file_path)
        suffix = file_path.suffix.lower()
        sha = sha256_file(file_path)
        doc_id = stable_id("doc", str(file_path.resolve()), sha, length=16)
        blocks: list[ParsedBlock] = []

        if self.config.prefer_docling and suffix in {".pdf", ".docx", ".xlsx", ".xlsm"}:
            docling_blocks = self._try_docling(file_path, doc_id)
            if docling_blocks:
                blocks.extend(docling_blocks)
        if not blocks:
            if suffix in {".txt", ".md", ".markdown"}:
                blocks.extend(self._parse_text_file(file_path, doc_id))
            elif suffix == ".csv":
                blocks.extend(self._parse_csv(file_path, doc_id))
            elif suffix in {".json", ".jsonl"}:
                blocks.extend(self._parse_jsonish(file_path, doc_id))
            elif suffix == ".pdf":
                blocks.extend(self._parse_pdf(file_path, doc_id))
            elif suffix == ".docx":
                blocks.extend(self._parse_docx(file_path, doc_id))
            elif suffix in {".xlsx", ".xlsm"}:
                blocks.extend(self._parse_xlsx(file_path, doc_id))
            elif suffix in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp", ".svg"}:
                blocks.extend(self._parse_image(file_path, doc_id))
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")

        return ParsedDocument(
            doc_id=doc_id,
            path=str(file_path),
            sha256=sha,
            parser_version=PARSER_VERSION,
            blocks=blocks,
            metadata={"suffix": suffix, "block_count": len(blocks)},
        )

    def _try_docling(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        try:
            from docling.document_converter import DocumentConverter  # type: ignore
        except Exception:
            return []
        try:
            converter = DocumentConverter()
            result = converter.convert(str(path))
            markdown = result.document.export_to_markdown()
        except Exception:
            return []
        return self._blocks_from_text(
            markdown,
            doc_id=doc_id,
            path=path,
            source_base=SourceRef(doc_id=doc_id, path=str(path)),
            origin="docling",
        )

    def _parse_text_file(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        text = read_text_lossy(path)
        return self._blocks_from_text(
            text,
            doc_id=doc_id,
            path=path,
            source_base=SourceRef(doc_id=doc_id, path=str(path)),
            origin="text",
        )

    def _parse_jsonish(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        if path.suffix.lower() == ".jsonl":
            lines = []
            for line in read_text_lossy(path).splitlines():
                if line.strip():
                    try:
                        lines.append(json.dumps(json.loads(line), ensure_ascii=False, indent=2))
                    except json.JSONDecodeError:
                        lines.append(line)
            text = "\n\n".join(lines)
        else:
            try:
                text = json.dumps(json.loads(read_text_lossy(path)), ensure_ascii=False, indent=2)
            except json.JSONDecodeError:
                text = read_text_lossy(path)
        return self._blocks_from_text(
            text,
            doc_id=doc_id,
            path=path,
            source_base=SourceRef(doc_id=doc_id, path=str(path)),
            origin="json",
        )

    def _parse_csv(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        rows: list[list[str]] = []
        with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as fh:
            reader = csv.reader(fh)
            rows.extend([str(cell) for cell in row] for row in reader)
        source = SourceRef(doc_id=doc_id, path=str(path), cell_range=csv_range(rows), table_no="table:1")
        return self._make_table_blocks(doc_id, rows, source, origin="csv", table_index=1)

    def _parse_image(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        payload = path.read_bytes()
        width, height = image_size(payload)
        media_type = guess_media_type(path)
        visual_type = classify_visual_type(file_name=path.name, media_type=media_type)
        ocr_text, ocr_meta = try_ocr_image(path) if self.config.enable_ocr else (None, {})
        source = SourceRef(doc_id=doc_id, path=str(path), asset_path=str(path), visual_no="image:1")
        blocks = [
            self._make_visual_block(
                doc_id=doc_id,
                source=source,
                visual_type=visual_type,
                caption=path.stem.replace("_", " "),
                nearby_text_value=ocr_text,
                asset_path=str(path),
                media_type=media_type,
                width=width,
                height=height,
                origin="image_file",
                visual_index=1,
                metadata_extra=ocr_meta,
            )
        ]
        if ocr_text:
            blocks.extend(
                self._blocks_from_text(
                    ocr_text,
                    doc_id=doc_id,
                    path=path,
                    source_base=SourceRef(doc_id=doc_id, path=str(path), asset_path=str(path), visual_no="image:1"),
                    origin="ocr",
                )
            )
        return blocks

    def _parse_pdf(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(str(path)) as pdf:
                for page_index, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    source = SourceRef(doc_id=doc_id, path=str(path), page=page_index)
                    blocks.extend(
                        self._blocks_from_text(
                            text,
                            doc_id=doc_id,
                            path=path,
                            source_base=source,
                            origin="pdfplumber",
                        )
                    )
                    for table_index, table in enumerate(page.extract_tables() or [], start=1):
                        table_source = SourceRef(
                            doc_id=doc_id,
                            path=str(path),
                            page=page_index,
                            section=f"table:{table_index}",
                            table_no=f"page:{page_index}.table:{table_index}",
                        )
                        blocks.extend(
                            self._make_table_blocks(
                                doc_id,
                                table,
                                table_source,
                                origin="pdfplumber",
                                table_index=table_index,
                            )
                        )
                    blocks.extend(self._visual_blocks_from_pdf_page(path, doc_id, page_index, page, text))
        except Exception:
            blocks.extend(self._parse_pdf_with_pypdf(path, doc_id))
        return blocks

    def _parse_pdf_with_pypdf(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception as exc:
            raise RuntimeError("No PDF parser available. Install pdfplumber or pypdf.") from exc
        reader = PdfReader(str(path))
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            source = SourceRef(doc_id=doc_id, path=str(path), page=page_index)
            blocks.extend(
                self._blocks_from_text(
                    text,
                    doc_id=doc_id,
                    path=path,
                    source_base=source,
                    origin="pypdf",
                )
            )
        return blocks

    def _parse_docx(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            raise RuntimeError("python-docx is required for DOCX parsing.") from exc

        document = Document(str(path))
        blocks: list[ParsedBlock] = []
        current_section: str | None = None
        paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs]
        paragraph_block_ids: dict[int, list[str]] = {}
        paragraph_sections: dict[int, str | None] = {}
        for idx, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = getattr(paragraph.style, "name", "") if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                current_section = text
            paragraph_sections[idx] = current_section
            source = SourceRef(doc_id=doc_id, path=str(path), section=current_section or f"p:{idx}")
            created = self._blocks_from_text(
                text,
                doc_id=doc_id,
                path=path,
                source_base=source,
                origin="docx",
                metadata_extra={"paragraph_index": idx, "style": style_name},
            )
            blocks.extend(created)
            paragraph_block_ids[idx] = [block.block_id for block in created if block.kind == "text"]
        for idx, paragraph in enumerate(document.paragraphs, start=1):
            blocks.extend(
                self._visual_blocks_from_docx_paragraph(
                    path=path,
                    doc_id=doc_id,
                    paragraph=paragraph,
                    paragraph_index=idx,
                    paragraph_texts=paragraph_texts,
                    linked_text_blocks=paragraph_block_ids.get(idx, []),
                    section=paragraph_sections.get(idx),
                )
            )
        for table_index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            source = SourceRef(
                doc_id=doc_id,
                path=str(path),
                section=f"table:{table_index}",
                table_no=f"table:{table_index}",
            )
            blocks.extend(
                self._make_table_blocks(
                    doc_id,
                    rows,
                    source,
                    origin="docx",
                    table_index=table_index,
                )
            )
        return blocks

    def _parse_xlsx(self, path: Path, doc_id: str) -> list[ParsedBlock]:
        try:
            import openpyxl  # type: ignore
            from openpyxl.utils import get_column_letter  # type: ignore
        except Exception as exc:
            raise RuntimeError("openpyxl is required for XLSX parsing.") from exc

        workbook = openpyxl.load_workbook(str(path), data_only=False, read_only=False)
        blocks: list[ParsedBlock] = []
        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            formula_cells: list[tuple[str, str]] = []
            for row in sheet.iter_rows():
                row_values: list[str] = []
                for cell in row:
                    value = cell.value
                    text = "" if value is None else str(value)
                    row_values.append(text)
                    if text.startswith("="):
                        formula_cells.append((cell.coordinate, text))
                if any(cell.strip() for cell in row_values):
                    rows.append(row_values)
            if rows:
                max_row = len(rows)
                max_col = max(len(row) for row in rows)
                cell_range = f"A1:{get_column_letter(max_col)}{max_row}"
                source = SourceRef(
                    doc_id=doc_id,
                    path=str(path),
                    sheet=sheet.title,
                    cell_range=cell_range,
                    table_no=f"sheet:{sheet.title}.table:1",
                )
                blocks.extend(
                    self._make_table_blocks(
                        doc_id,
                        rows,
                        source,
                        origin="xlsx",
                        table_index=1,
                        sheet=sheet.title,
                    )
                )
            for coordinate, formula in formula_cells:
                source = SourceRef(doc_id=doc_id, path=str(path), sheet=sheet.title, cell_range=coordinate)
                blocks.append(
                    self._make_block(
                        doc_id,
                        "cell",
                        formula,
                        source,
                        metadata={"origin": "xlsx", "sheet": sheet.title, "cell": coordinate, "is_formula": True},
                    )
                )
            blocks.extend(self._visual_blocks_from_xlsx_sheet(path, doc_id, sheet))
        return blocks

    def _make_table_blocks(
        self,
        doc_id: str,
        rows: Iterable[Iterable[object]],
        source: SourceRef,
        *,
        origin: str,
        table_index: int,
        sheet: str | None = None,
    ) -> list[ParsedBlock]:
        normalized_rows = normalize_rows(rows)
        if not normalized_rows:
            return []
        table_text = rows_to_tsv(normalized_rows)
        table_no = source.table_no or f"table:{table_index}"
        table_source = source.model_copy(update={"table_no": table_no})
        table_block = self._make_block(
            doc_id,
            "table",
            table_text,
            table_source,
            metadata={
                "origin": origin,
                "table_index": table_index,
                "sheet": sheet,
                "row_count": len(normalized_rows),
                "column_count": max(len(row) for row in normalized_rows),
            },
        )
        blocks = [table_block]
        header = infer_header(normalized_rows)
        for row_index, row in enumerate(normalized_rows, start=1):
            if not any(cell.strip() for cell in row):
                continue
            label = row_label(row, header)
            values = row_values(row, header)
            row_text = build_table_row_text(label=label, values=values, raw=row)
            row_source = table_source.model_copy(update={"row_index": row_index})
            row_block = self._make_block(
                doc_id,
                "table_row",
                row_text,
                row_source,
                metadata={
                    "origin": origin,
                    "table_index": table_index,
                    "parent_block_id": table_block.block_id,
                    "row_index": row_index,
                    "row_label": label,
                    "cells": row,
                    "values": values,
                },
            )
            blocks.append(row_block)
            for col_index, cell_text in enumerate(row, start=1):
                if not cell_text.strip():
                    continue
                column_name = header[col_index - 1] if header and col_index - 1 < len(header) else None
                cell_source = table_source.model_copy(update={"row_index": row_index, "col_index": col_index})
                blocks.append(
                    self._make_block(
                        doc_id,
                        "table_cell",
                        build_table_cell_text(
                            row_label=label,
                            column_name=column_name,
                            cell_text=cell_text,
                            row_index=row_index,
                            col_index=col_index,
                        ),
                        cell_source,
                        metadata={
                            "origin": origin,
                            "table_index": table_index,
                            "parent_block_id": table_block.block_id,
                            "row_block_id": row_block.block_id,
                            "row_index": row_index,
                            "col_index": col_index,
                            "row_label": label,
                            "column_name": column_name,
                            "cell_text": cell_text,
                        },
                    )
                )
        return blocks

    def _blocks_from_text(
        self,
        text: str,
        *,
        doc_id: str,
        path: Path,
        source_base: SourceRef,
        origin: str,
        metadata_extra: dict | None = None,
    ) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        if not text.strip():
            return blocks

        chunks = split_text(text, self.config.chunk_target_chars, self.config.chunk_overlap_chars)
        for chunk_index, chunk in enumerate(chunks):
            section = source_base.section or infer_section(chunk)
            source = source_base.model_copy(update={"section": section})
            text_block = self._make_block(
                doc_id,
                "text",
                chunk,
                source,
                metadata={**(metadata_extra or {}), "origin": origin, "chunk_index": chunk_index},
            )
            blocks.append(text_block)
            for visual_index, visual in enumerate(extract_markdown_visuals(chunk), start=1):
                asset_path = copy_local_visual(
                    source_doc=path,
                    image_ref=visual.src,
                    artifact_dir=Path(self.config.artifact_dir),
                    doc_id=doc_id,
                )
                visual_source = source_base.model_copy(
                    update={
                        "section": section,
                        "asset_path": asset_path,
                        "visual_no": f"markdown:{chunk_index + 1}.{visual_index}",
                    }
                )
                blocks.append(
                    self._make_visual_block(
                        doc_id=doc_id,
                        source=visual_source,
                        visual_type=visual.visual_type,
                        caption=visual.caption,
                        nearby_text_value=visual.nearby_text,
                        asset_path=asset_path,
                        media_type=guess_media_type(Path(visual.src)),
                        origin=origin,
                        visual_index=visual_index,
                        linked_text_blocks=[text_block.block_id],
                        metadata_extra={
                            "markdown_src": visual.src,
                            "char_span": [visual.start, visual.end],
                        },
                    )
                )
            if self.config.enable_formula_detection:
                for formula_index, candidate in enumerate(extract_formula_candidates(chunk), start=1):
                    equation_source = source.model_copy(
                        update={
                            "equation_no": candidate.equation_no,
                            "section": section,
                        }
                    )
                    blocks.append(
                        self._make_block(
                            doc_id,
                            "equation",
                            candidate.latex,
                            equation_source,
                            latex=candidate.latex,
                            normalized_latex=candidate.normalized_latex,
                            symbols=candidate.symbols,
                            operators=candidate.operators,
                            metadata={
                                "origin": origin,
                                "parent_block_id": text_block.block_id,
                                "formula_index": formula_index,
                                "char_span": [candidate.start, candidate.end],
                                "context": trim(chunk, 700),
                            },
                        )
                    )
        return blocks

    def _visual_blocks_from_pdf_page(self, path: Path, doc_id: str, page_index: int, page, text: str) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        lines = text.splitlines()
        caption = nearest_caption(lines)
        for image_index, image in enumerate(getattr(page, "images", []) or [], start=1):
            bbox = [float(image.get(key, 0.0) or 0.0) for key in ("x0", "top", "x1", "bottom")]
            width = int(abs((image.get("x1") or 0) - (image.get("x0") or 0))) or None
            height = int(abs((image.get("bottom") or 0) - (image.get("top") or 0))) or None
            visual_type = classify_visual_type(caption=caption, nearby_text=text, explicit_hint="pdf_image")
            source = SourceRef(
                doc_id=doc_id,
                path=str(path),
                page=page_index,
                bbox=bbox,
                visual_no=f"page:{page_index}.image:{image_index}",
            )
            blocks.append(
                self._make_visual_block(
                    doc_id=doc_id,
                    source=source,
                    visual_type=visual_type,
                    caption=caption,
                    nearby_text_value=compact_text(text, 1000),
                    asset_path=None,
                    media_type=None,
                    width=width,
                    height=height,
                    origin="pdfplumber_image",
                    visual_index=image_index,
                    metadata_extra={"raw_image": compact_mapping(image)},
                )
            )
        if not blocks:
            for caption_index, line in enumerate(lines, start=1):
                if not is_caption_line(line):
                    continue
                visual_type = classify_visual_type(caption=line, nearby_text=text)
                source = SourceRef(
                    doc_id=doc_id,
                    path=str(path),
                    page=page_index,
                    visual_no=f"page:{page_index}.caption:{caption_index}",
                )
                blocks.append(
                    self._make_visual_block(
                        doc_id=doc_id,
                        source=source,
                        visual_type=visual_type,
                        caption=line.strip(),
                        nearby_text_value=compact_text(text, 1000),
                        asset_path=None,
                        media_type=None,
                        origin="pdf_caption",
                        visual_index=caption_index,
                    )
                )
        return blocks

    def _visual_blocks_from_docx_paragraph(
        self,
        *,
        path: Path,
        doc_id: str,
        paragraph,
        paragraph_index: int,
        paragraph_texts: list[str],
        linked_text_blocks: list[str],
        section: str | None,
    ) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        nearby = neighbor_text(paragraph_texts, paragraph_index - 1)
        caption = caption_for_position(paragraph_texts, paragraph_index - 1)
        paragraph_text = paragraph_texts[paragraph_index - 1] if paragraph_index - 1 < len(paragraph_texts) else ""
        if paragraph_text and is_caption_line(paragraph_text):
            caption = paragraph_text

        blips = paragraph._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        for image_index, blip in enumerate(blips, start=1):
            rel_id = (
                blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                or blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link")
            )
            part = paragraph.part.related_parts.get(rel_id) if rel_id else None
            payload = getattr(part, "blob", b"") if part is not None else b""
            media_type = getattr(part, "content_type", None) if part is not None else None
            width, height = image_size(payload) if payload else (None, None)
            asset_path = (
                save_visual_artifact(
                    artifact_dir=Path(self.config.artifact_dir),
                    doc_id=doc_id,
                    payload=payload,
                    preferred_name=f"docx_p{paragraph_index}_image_{image_index}",
                    media_type=media_type,
                )
                if payload
                else None
            )
            visual_type = classify_visual_type(caption=caption, nearby_text=nearby, media_type=media_type)
            source = SourceRef(
                doc_id=doc_id,
                path=str(path),
                section=section or f"p:{paragraph_index}",
                asset_path=asset_path,
                visual_no=f"paragraph:{paragraph_index}.image:{image_index}",
            )
            blocks.append(
                self._make_visual_block(
                    doc_id=doc_id,
                    source=source,
                    visual_type=visual_type,
                    caption=caption,
                    nearby_text_value=nearby,
                    asset_path=asset_path,
                    media_type=media_type,
                    width=width,
                    height=height,
                    origin="docx_image",
                    visual_index=image_index,
                    linked_text_blocks=linked_text_blocks,
                    metadata_extra={"rel_id": rel_id},
                )
            )

        charts = paragraph._element.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/chart}chart")
        for chart_index, chart in enumerate(charts, start=1):
            rel_id = chart.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            source = SourceRef(
                doc_id=doc_id,
                path=str(path),
                section=section or f"p:{paragraph_index}",
                visual_no=f"paragraph:{paragraph_index}.chart:{chart_index}",
            )
            blocks.append(
                self._make_visual_block(
                    doc_id=doc_id,
                    source=source,
                    visual_type="chart",
                    caption=caption,
                    nearby_text_value=nearby,
                    asset_path=None,
                    media_type="application/vnd.openxmlformats-officedocument.drawingml.chart+xml",
                    origin="docx_chart",
                    visual_index=chart_index,
                    linked_text_blocks=linked_text_blocks,
                    metadata_extra={"rel_id": rel_id},
                )
            )
        return blocks

    def _visual_blocks_from_xlsx_sheet(self, path: Path, doc_id: str, sheet) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        for chart_index, chart in enumerate(getattr(sheet, "_charts", []) or [], start=1):
            cell = anchor_cell(getattr(chart, "anchor", None))
            title = chart_title(chart)
            nearby = nearby_sheet_text(sheet, cell)
            caption = title or nearest_caption(nearby.splitlines())
            source = SourceRef(
                doc_id=doc_id,
                path=str(path),
                sheet=sheet.title,
                cell_range=cell,
                visual_no=f"sheet:{sheet.title}.chart:{chart_index}",
            )
            blocks.append(
                self._make_visual_block(
                    doc_id=doc_id,
                    source=source,
                    visual_type="chart",
                    caption=caption,
                    nearby_text_value=nearby,
                    asset_path=None,
                    media_type="application/vnd.openxmlformats-officedocument.drawingml.chart+xml",
                    origin="xlsx_chart",
                    visual_index=chart_index,
                    metadata_extra={"chart_class": chart.__class__.__name__},
                )
            )
        for image_index, image in enumerate(getattr(sheet, "_images", []) or [], start=1):
            cell = anchor_cell(getattr(image, "anchor", None))
            nearby = nearby_sheet_text(sheet, cell)
            caption = nearest_caption(nearby.splitlines())
            payload = image._data() if hasattr(image, "_data") else b""
            media_type = guess_media_type(Path(getattr(image, "path", "") or "image"))
            width = getattr(image, "width", None)
            height = getattr(image, "height", None)
            asset_path = (
                save_visual_artifact(
                    artifact_dir=Path(self.config.artifact_dir),
                    doc_id=doc_id,
                    payload=payload,
                    preferred_name=f"xlsx_{sheet.title}_image_{image_index}",
                    media_type=media_type,
                )
                if payload
                else None
            )
            visual_type = classify_visual_type(caption=caption, nearby_text=nearby, media_type=media_type)
            source = SourceRef(
                doc_id=doc_id,
                path=str(path),
                sheet=sheet.title,
                cell_range=cell,
                asset_path=asset_path,
                visual_no=f"sheet:{sheet.title}.image:{image_index}",
            )
            blocks.append(
                self._make_visual_block(
                    doc_id=doc_id,
                    source=source,
                    visual_type=visual_type,
                    caption=caption,
                    nearby_text_value=nearby,
                    asset_path=asset_path,
                    media_type=media_type,
                    width=width,
                    height=height,
                    origin="xlsx_image",
                    visual_index=image_index,
                )
            )
        return blocks

    def _make_visual_block(
        self,
        *,
        doc_id: str,
        source: SourceRef,
        visual_type: str,
        caption: str | None,
        nearby_text_value: str | None,
        asset_path: str | None = None,
        media_type: str | None = None,
        width: int | None = None,
        height: int | None = None,
        origin: str,
        visual_index: int,
        linked_text_blocks: list[str] | None = None,
        metadata_extra: dict | None = None,
    ) -> ParsedBlock:
        label = source.visual_no or f"visual:{visual_index}"
        text = build_visual_index_text(
            visual_type=visual_type,
            caption=caption,
            nearby_text_value=nearby_text_value,
            label=label,
            file_name=asset_path,
        )
        metadata = {
            "origin": origin,
            "visual_type": visual_type,
            "caption": caption,
            "nearby_text": nearby_text_value,
            "asset_path": asset_path,
            "media_type": media_type,
            "width": width,
            "height": height,
            "linked_text_blocks": linked_text_blocks or [],
            "visual_index": visual_index,
        }
        if metadata_extra:
            metadata.update(metadata_extra)
        return self._make_block(
            doc_id,
            "visual_object",
            text,
            source,
            metadata=metadata,
        )

    def _make_block(
        self,
        doc_id: str,
        kind: str,
        text: str,
        source: SourceRef,
        *,
        latex: str | None = None,
        normalized_latex: str | None = None,
        symbols: list[str] | None = None,
        operators: list[str] | None = None,
        metadata: dict | None = None,
    ) -> ParsedBlock:
        metadata_payload = json.dumps(metadata or {}, ensure_ascii=False, sort_keys=True, default=str)
        block_id = stable_id("blk", doc_id, kind, source.model_dump_json(), metadata_payload, text[:500], length=18)
        final_source = source.model_copy(update={"block_id": block_id})
        return ParsedBlock(
            block_id=block_id,
            doc_id=doc_id,
            kind=kind,  # type: ignore[arg-type]
            text=text.strip(),
            source=final_source,
            latex=latex,
            normalized_latex=normalized_latex,
            symbols=symbols or [],
            operators=operators or [],
            metadata=metadata or {},
        )


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_text_lossy(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def rows_to_tsv(rows: Iterable[Iterable[object]]) -> str:
    lines = []
    for row in rows:
        values = ["" if cell is None else str(cell).replace("\n", " ").strip() for cell in row]
        lines.append("\t".join(values).rstrip())
    return "\n".join(line for line in lines if line.strip())


def normalize_rows(rows: Iterable[Iterable[object]]) -> list[list[str]]:
    normalized: list[list[str]] = []
    max_cols = 0
    for row in rows:
        values = ["" if cell is None else str(cell).replace("\r\n", "\n").replace("\r", "\n").strip() for cell in row]
        if not any(values):
            continue
        max_cols = max(max_cols, len(values))
        normalized.append(values)
    if max_cols <= 0:
        return []
    return [row + [""] * (max_cols - len(row)) for row in normalized]


def infer_header(rows: list[list[str]]) -> list[str] | None:
    if not rows:
        return None
    first = [cell.strip() for cell in rows[0]]
    non_empty = [cell for cell in first if cell]
    if len(non_empty) >= 2 and all(len(cell) <= 40 for cell in non_empty):
        lower = {cell.lower() for cell in non_empty}
        if len(lower & {"id", "name", "description", "requirement", "value", "parameter", "type"}) >= 2:
            return first
        if len(non_empty) >= 3:
            return first
    return None


def row_label(row: list[str], header: list[str] | None = None) -> str | None:
    if header and len(row) >= 2:
        for idx, column in enumerate(header):
            if column.strip().lower() in {"id", "name", "field", "key", "type"} and idx < len(row) and row[idx].strip():
                return row[idx].strip()
    for cell in row:
        value = " ".join(cell.split())
        if value and len(value) <= 40:
            return value
    return None


def row_values(row: list[str], header: list[str] | None = None) -> dict[str, str]:
    values: dict[str, str] = {}
    if header:
        for idx, cell in enumerate(row):
            key = header[idx].strip() if idx < len(header) and header[idx].strip() else f"col_{idx + 1}"
            if cell.strip():
                values[key] = cell.strip()
        return values
    if len(row) >= 2:
        first = row[0].strip()
        rest = "\n".join(cell.strip() for cell in row[1:] if cell.strip())
        if first and rest:
            return {first: rest}
    for idx, cell in enumerate(row, start=1):
        if cell.strip():
            values[f"col_{idx}"] = cell.strip()
    return values


def build_table_row_text(*, label: str | None, values: dict[str, str], raw: list[str]) -> str:
    parts = ["table_row"]
    if label:
        parts.append(f"row_label: {label}")
    for key, value in values.items():
        parts.append(f"{key}: {value}")
    if not values:
        parts.append("raw: " + " | ".join(cell for cell in raw if cell.strip()))
    return "\n".join(parts)


def build_table_cell_text(
    *,
    row_label: str | None,
    column_name: str | None,
    cell_text: str,
    row_index: int,
    col_index: int,
) -> str:
    parts = ["table_cell", f"row:{row_index}", f"col:{col_index}"]
    if row_label:
        parts.append(f"row_label: {row_label}")
    if column_name:
        parts.append(f"column: {column_name}")
    parts.append(f"text: {cell_text}")
    return "\n".join(parts)


def csv_range(rows: list[list[str]]) -> str | None:
    if not rows:
        return None
    try:
        from openpyxl.utils import get_column_letter  # type: ignore

        max_col = max(len(row) for row in rows)
        return f"A1:{get_column_letter(max_col)}{len(rows)}"
    except Exception:
        return None


def split_text(text: str, target_chars: int = 1200, overlap_chars: int = 180) -> list[str]:
    clean = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [part.strip() for part in clean.split("\n\n") if part.strip()]
    if not paragraphs:
        paragraphs = [line.strip() for line in clean.splitlines() if line.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= target_chars:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
            current = tail(current, overlap_chars)
        if len(paragraph) > target_chars:
            for piece in hard_wrap(paragraph, target_chars, overlap_chars):
                if piece:
                    chunks.append(piece)
            current = ""
        else:
            current = f"{current}\n\n{paragraph}".strip()
    if current:
        chunks.append(current)
    return chunks or [clean.strip()]


def hard_wrap(text: str, target_chars: int, overlap_chars: int) -> list[str]:
    pieces: list[str] = []
    start = 0
    step = max(1, target_chars - overlap_chars)
    while start < len(text):
        pieces.append(text[start : start + target_chars].strip())
        start += step
    return pieces


def tail(text: str, chars: int) -> str:
    if chars <= 0:
        return ""
    return text[-chars:].strip()


def trim(text: str, chars: int) -> str:
    value = " ".join(text.split())
    if len(value) <= chars:
        return value
    return value[: chars - 3] + "..."


def infer_section(text: str) -> str | None:
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            return stripped.lstrip("#").strip()
        if len(stripped) < 100 and (stripped[:1].isdigit() or stripped.endswith(":")):
            return stripped.rstrip(":")
        return None
    return None


def guess_media_type(path: Path) -> str | None:
    suffix = path.suffix.lower()
    mapping = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
        ".webp": "image/webp",
        ".svg": "image/svg+xml",
    }
    return mapping.get(suffix)


def compact_mapping(value: dict) -> dict:
    compact: dict = {}
    for key, item in value.items():
        if isinstance(item, (str, int, float, bool)) or item is None:
            compact[str(key)] = item
    return compact


def neighbor_text(lines: list[str], index: int, radius: int = 2) -> str:
    left = max(0, index - radius)
    right = min(len(lines), index + radius + 1)
    return compact_text("\n".join(line for line in lines[left:right] if line.strip()), 1200)


def caption_for_position(lines: list[str], index: int) -> str | None:
    if 0 <= index < len(lines) and is_caption_line(lines[index]):
        return lines[index].strip()
    return nearest_caption(lines[max(0, index - 2) : min(len(lines), index + 3)])


def anchor_cell(anchor) -> str | None:
    if isinstance(anchor, str):
        return anchor
    marker = getattr(anchor, "_from", None)
    if marker is None:
        return None
    try:
        from openpyxl.utils import get_column_letter  # type: ignore

        return f"{get_column_letter(int(marker.col) + 1)}{int(marker.row) + 1}"
    except Exception:
        return None


def nearby_sheet_text(sheet, cell: str | None, radius: int = 3) -> str:
    if not cell:
        return ""
    try:
        from openpyxl.utils.cell import coordinate_to_tuple  # type: ignore

        row, col = coordinate_to_tuple(cell)
    except Exception:
        return ""
    lines: list[str] = []
    for row_idx in range(max(1, row - radius), min(sheet.max_row, row + radius) + 1):
        values: list[str] = []
        for col_idx in range(max(1, col - radius), min(sheet.max_column, col + radius) + 1):
            value = sheet.cell(row=row_idx, column=col_idx).value
            if value is not None:
                values.append(str(value))
        if values:
            lines.append("\t".join(values))
    return compact_text("\n".join(lines), 1200)


def chart_title(chart) -> str | None:
    title = getattr(chart, "title", None)
    if title is None:
        return None
    text = str(title)
    if len(text) > 180 or "openpyxl" in text.lower():
        return None
    return text.strip() or None
